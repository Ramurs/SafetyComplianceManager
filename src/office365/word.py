from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import get_settings


def generate_audit_report(audit) -> Path:
    """Generate a Word document for an audit report."""
    settings = get_settings()
    doc = Document()

    # Title
    title = doc.add_heading(audit.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Status: {audit.status}")
    doc.add_paragraph(f"Date: {audit.created_at.strftime('%Y-%m-%d')}")
    if audit.scope:
        doc.add_paragraph(f"Scope: {audit.scope}")

    # Summary
    if audit.summary:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(audit.summary)

    # Findings
    if audit.findings:
        doc.add_heading("Findings", level=1)

        severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
        sorted_findings = sorted(audit.findings, key=lambda f: severity_order.get(f.severity, 5))

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "Severity"
        headers[1].text = "Control"
        headers[2].text = "Finding"
        headers[3].text = "Recommendation"

        for finding in sorted_findings:
            row = table.add_row().cells
            row[0].text = finding.severity.upper()
            row[1].text = finding.control_id or ""
            row[2].text = f"{finding.title}\n{finding.description}"
            row[3].text = finding.recommendation

    output_path = settings.output_dir / f"audit_{audit.id[:8]}.docx"
    doc.save(str(output_path))
    return output_path


def generate_policy_document(policy) -> Path:
    """Generate a Word document for a policy."""
    settings = get_settings()
    doc = Document()

    title = doc.add_heading(policy.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Version: {policy.current_version}")
    doc.add_paragraph(f"Status: {policy.status}")
    doc.add_paragraph(f"Category: {policy.category}")
    doc.add_paragraph(f"Last Updated: {policy.updated_at.strftime('%Y-%m-%d')}")

    doc.add_heading("Policy Content", level=1)

    if policy.versions:
        latest = max(policy.versions, key=lambda v: v.version_number)
        for paragraph in latest.content.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    output_path = settings.output_dir / f"policy_{policy.id[:8]}.docx"
    doc.save(str(output_path))
    return output_path
